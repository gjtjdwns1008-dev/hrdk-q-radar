# -*- coding: utf-8 -*-
"""
briefing_maker.py
=================
[월별 이슈브리핑 + 모니터링 결과 자동 생성기]

GitHub Actions에서 월(예: 202601)을 입력받아 실행됩니다.

전체 흐름:
  1) 구글 시트에서 해당 월 데이터 호출 (연관높음 + 단순관련)
  2) Gemini가 ① TOP N 선별 → ② 총평 작성 → ③ 법령별 상세 작성
  3) 이슈브리핑(.docx) + 모니터링 결과(.xlsx) 생성  ← 우리가 디자인한 그대로
  4) 두 파일을 웹훅(Make.com)으로 첨부 발송 → 공단 메일함

[필요 환경변수] (GitHub Secrets에 저장)
  GCP_SA_JSON      : 구글 서비스계정 인증 JSON
  GOOGLE_SHEET_ID  : monitor 구글시트 KEY
  GEMINI_API_KEY   : Gemini API 키
  LLM_MODEL        : (선택) 모델명. 기본 gemini-2.5-pro
  BRIEFING_WEBHOOK_URL : (선택) 이슈브리핑 전용 Make.com 웹훅. 비우면 메일 미발송(로컬 테스트용)
                         ※ monitor 일일 알림 웹훅(WEBHOOK_URL)과는 별개의 새 웹훅
  TARGET_MONTH     : 생성할 월 (예: 202601). 워크플로우 입력값으로 주입.
"""

import os
import io
import re
import json
import time
import calendar
import requests
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager
from collections import Counter

# 구글 시트 인증은 기존 코어 모듈을 그대로 재사용
from hrdk_law_core.sheets import get_sheet_client
# Gemini 호출도 기존 모델 추상화 모듈 재사용 (모델 교체 가능)
from hrdk_law_core.llm_client import get_llm_client
from hrdk_law_core.certs import TRACK1_TYPE_KO, TRACK1_RISK_KO, TRACK2_CODE_KO

# 프롬프트는 별도 파일에서
from briefing_prompts import (
    PERSONA, SELECT_TOP_PROMPT, FOREWORD_PROMPT, DETAIL_PROMPT
)

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from openpyxl import Workbook
from openpyxl.styles import Font as XLFont, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ============================================================
# 설정값
# ============================================================
def _last_month():
    """오늘 기준 '지난달'을 YYYYMM으로 반환. (자동 실행 시 사용)
    매달 2일에 돌면 지난달 보고서를 만든다. 연초(1월)엔 작년 12월로 넘어간다.
    예: 오늘이 2026-02-02 → '202601' / 2026-01-02 → '202512'
    """
    from datetime import datetime, timezone, timedelta
    # GitHub Actions는 UTC로 돈다. 한국 날짜 기준으로 '지난달'을 정확히 잡기 위해
    # KST(UTC+9)로 변환한 뒤 계산한다.
    now_kst = datetime.now(timezone.utc) + timedelta(hours=9)
    year, month = now_kst.year, now_kst.month
    if month == 1:
        return f"{year - 1}12"
    return f"{year}{month - 1:02d}"


# TARGET_MONTH: 수동 입력값이 있으면 그걸 쓰고, 없으면(자동 실행) 지난달.
TARGET_MONTH = os.environ.get("TARGET_MONTH", "").strip()  # 예: 202601
if not TARGET_MONTH:
    TARGET_MONTH = _last_month()
    print(f"ℹ️ TARGET_MONTH 미입력 → 자동으로 지난달({TARGET_MONTH}) 생성")
GCP_SA_JSON = os.environ.get("GCP_SERVICE_ACCOUNT_JSON") or os.environ.get("GCP_SA_JSON")
GOOGLE_SHEET_ID = os.environ.get("GOOGLE_SHEET_URL") or os.environ.get("GOOGLE_SHEET_ID")
WEBHOOK_URL = os.environ.get("BRIEFING_WEBHOOK_URL")  # 이슈브리핑 전용 웹훅 (monitor 일일 알림과 분리)

TOP_N = 5  # 이슈브리핑에 담을 핵심 법령 수 (기본 5건)

# 디자인 색상 (우리가 정한 공단 네이비 톤)
NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x5A, 0x88)
GRAY = RGBColor(0x59, 0x59, 0x59)
DARKGRAY = RGBColor(0x40, 0x40, 0x40)
FONT = "맑은 고딕"

MAIN_TAB = "국가기술자격 관련법령"  # Q-RADAR 통합 대장 — 연관도 컬럼으로 분리


# ============================================================
# [1] 구글 시트에서 해당 월 데이터 호출
# ============================================================
def norm_date(v):
    s = "".join(c for c in str(v) if c.isdigit())
    return s[:8] if len(s) >= 8 else ""


def fetch_month_data(target_month):
    """구글 시트에서 target_month(YYYYMM) 데이터를 가져온다.
    반환: (high, simple, total_laws)
      - high   : 활용 높은 법령 목록
      - simple : 단순 관련 법령 목록
      - total_laws : 그달 전체 시행 법령 수 (총괄현황표의 '총 검토건수' 합계)

    ※ 법제처 API를 따로 부르지 않는다.
      monitor가 매일 검토하며 총괄현황표에 쌓아둔 '총 검토건수'가 곧 전체 시행 법령 수다.
      (구글 시트가 우리의 영구 스토리지이므로, 모든 통계를 여기서 가져온다.)
    """
    print(f"📥 [1단계] 구글 시트에서 {target_month} 데이터 호출 중...")
    # get_sheet_client는 (client, spreadsheet) 튜플 반환 → 두 번째만 사용
    _, ss = get_sheet_client(GCP_SA_JSON, GOOGLE_SHEET_ID)

    records = ss.worksheet(MAIN_TAB).get_all_records()
    month_rows = [r for r in records
                  if norm_date(r.get("시행일자", "")).startswith(target_month)]
    high = [r for r in month_rows if str(r.get("연관도", "")).strip() == "연관높음"]
    simple = [r for r in month_rows if str(r.get("연관도", "")).strip() == "단순관련"]
    preferred = [r for r in month_rows if str(r.get("우대여부", "")).strip() == "O"]

    # 총괄현황표에서 그달 '총 검토건수' 합계 = 전체 시행 법령 수
    total_laws = sum_total_reviewed(ss, target_month)

    print(f"   → 전체 {total_laws if total_laws else '-(미집계)'}건 / 연관높음 {len(high)} / 단순관련 {len(simple)} / 우대 {len(preferred)}")
    return high, simple, preferred, total_laws


def sum_total_reviewed(ss, target_month):
    """총괄현황표에서 target_month(YYYYMM)에 해당하는 '총 검토건수'를 모두 더한다.
    이 값이 곧 그달 전체 시행(검토 대상) 법령 수다. 실패 시 None."""
    try:
        records = ss.worksheet("총괄현황표").get_all_records()
    except Exception as e:
        print(f"   ⚠️ 총괄현황표 읽기 실패: {e}")
        return None

    # '총 검토건수' 칸 이름이 다를 수 있어 후보를 둔다 (헤더 표기 차이 대응)
    count_keys = ["총 검토건수", "총검토건수", "총 검토 건수"]
    total = 0
    found = False
    for r in records:
        # 날짜 칸도 표기 차이 대응 (시행일자/수집일자)
        date_val = r.get("시행일자") or r.get("수집일자") or ""
        if not norm_date(date_val).startswith(target_month):
            continue
        # 총 검토건수 칸 찾기
        for k in count_keys:
            if k in r:
                try:
                    total += int(float(r[k] or 0))
                    found = True
                except (ValueError, TypeError):
                    pass
                break
    return total if found else None


# ============================================================
# [2] Gemini 3단계 호출
# ============================================================
def _ask(prompt, temperature=0.2, retries=3):
    """Gemini 호출 + 재시도. 코어의 llm_client 사용 (웹검색 없음)."""
    llm = get_llm_client()
    for attempt in range(retries):
        try:
            return llm.generate(prompt, temperature=temperature)
        except Exception as e:
            print(f"   ⚠️ AI 호출 실패({attempt+1}/{retries}): {e}")
            time.sleep(10)
    return ""


# 웹검색(grounding) 사용 가능 여부 — 환경변수로 켜고 끔 (기본: 켜기)
USE_WEB_SEARCH = os.environ.get("USE_WEB_SEARCH", "true").lower() in ("true", "1", "yes")


def _ask_with_search(prompt, temperature=0.2):
    """웹검색(Google Search grounding)을 켜서 Gemini 호출.
    상세분석처럼 '법령의 실제 내용'을 풍부하게 써야 할 때 사용.

    - 코어(llm_client)는 건드리지 않고, 여기서 직접 Gemini를 호출한다.
      (코어는 monitor/RADAR 공용이라 검색 기능을 함부로 넣지 않음)
    - 검색이 안 되거나 실패하면 일반 호출(_ask)로 자동 폴백한다.
      → 검색이 되면 더 자세하게, 안 되면 최소한 시트 기반으로는 나온다.
    """
    if not USE_WEB_SEARCH:
        return _ask(prompt, temperature=temperature)

    api_key = os.environ.get("GEMINI_API_KEY")
    model = os.environ.get("LLM_MODEL", "") or "gemini-3.5-flash"
    if not api_key:
        return _ask(prompt, temperature=temperature)

    try:
        from google import genai
        from google.genai import types
        client = genai.Client(api_key=api_key)
        # Google Search 도구 활성화
        search_tool = types.Tool(google_search=types.GoogleSearch())
        resp = client.models.generate_content(
            model=model,
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=temperature,
                tools=[search_tool],
                max_output_tokens=8192,
            ),
        )
        text = (resp.text or "").strip()
        if text:
            return text
        # 빈 응답이면 폴백
        print("   ℹ️ 웹검색 응답이 비어 일반 호출로 폴백")
        return _ask(prompt, temperature=temperature)
    except Exception as e:
        # 검색 미지원 모델·쿼터·기타 오류 → 일반 호출로 폴백
        print(f"   ℹ️ 웹검색 사용 불가({str(e)[:40]}) → 일반 호출로 폴백")
        return _ask(prompt, temperature=temperature)


def _select_rank_key(r):
    """결정적 정렬: ① 대폭 감소 우선(정책 대응 시급) ② 관련 종목 수 많은 순(파급 범위) ③ 법령명"""
    certs = [c for c in str(r.get("관련 종목", "")).split(",") if c.strip() and c.strip() != "없음"]
    return (0 if r.get("활용도_구분", "") == "대폭 감소" else 1, -len(certs), str(r.get("법령명", "")))


def _auto_reason(r):
    n = len([c for c in str(r.get("관련 종목", "")).split(",") if c.strip() and c.strip() != "없음"])
    return f"{r.get('활용도_구분', '')} · 관련 종목 {n}개 파급"


def select_top_laws(big_laws, top_n=TOP_N):
    """[C안] 결정적 압축 → AI가 N건 + 선정 사유 → 실패 시 결정적 상위 N (앞N건 폴백 폐지)."""
    print(f"🧠 [2-1] 대폭 증감 {len(big_laws)}건 중 핵심 {top_n}건 선별 중... (C안: 압축+AI사유)")
    ranked = sorted(big_laws, key=_select_rank_key)
    if len(ranked) <= top_n:
        for r in ranked:
            r["_선정사유"] = r.get("_선정사유") or _auto_reason(r)
        return ranked
    pool = ranked[:max(top_n + 3, 8)]
    candidates = [{"id": i, "법령명": r.get("법령명", ""), "활용도": r.get("활용도_구분", ""),
                   "관련자격": str(r.get("관련 종목", ""))[:120],
                   "주요내용": str(r.get("주요 제·개정내용", ""))[:160]}
                  for i, r in enumerate(pool)]
    try:
        raw = _ask(SELECT_TOP_PROMPT.format(persona=PERSONA, top_n=top_n,
                                            candidates=json.dumps(candidates, ensure_ascii=False)),
                   temperature=0.0)
        m = re.search(r"\[.*\]", raw, re.DOTALL)
        items = json.loads(m.group(0)) if m else []
        picked = []
        for it in items:
            if isinstance(it, dict) and isinstance(it.get("id"), int) and 0 <= it["id"] < len(pool):
                r = pool[it["id"]]
                r["_선정사유"] = _clean_citations(str(it.get("reason", ""))).strip() or _auto_reason(r)
                if r not in picked:
                    picked.append(r)
            elif isinstance(it, int) and 0 <= it < len(pool):  # 구형 [1,2,…] 응답 호환
                r = pool[it]
                r["_선정사유"] = _auto_reason(r)
                if r not in picked:
                    picked.append(r)
        if picked:
            return picked[:top_n]
    except Exception:
        pass
    print("   ⚠️ AI 선별 실패 → 결정적 상위 N건 (감소 우선·종목 수 순)")
    for r in pool[:top_n]:
        r["_선정사유"] = r.get("_선정사유") or _auto_reason(r)
    return pool[:top_n]


def make_foreword(selected, target_month):
    """이달의 총평 작성."""
    print("🧠 [2-2] 이달의 총평 작성 중...")
    year, month = target_month[:4], str(int(target_month[4:6]))
    summary = [{"법령명": r.get("법령명", ""),
                "내용": r.get("주요 제·개정내용", "")} for r in selected]
    prompt = FOREWORD_PROMPT.format(
        persona=PERSONA, year=year, month=month, count=len(selected),
        summary_data=json.dumps(summary, ensure_ascii=False)
    )
    text = _ask(prompt, temperature=0.3).strip()
    if not text:
        text = (f"{year}년 {month}월은 국가기술자격의 활용 기반이 전반적으로 "
                f"강화되는 추세가 확인되는 시기입니다.")
    return text


def _code(v):
    """'B (영업요건형)' → 'B' — 라벨 병기 값에서 코드만 추출"""
    return str(v or "").strip().split(" ")[0].split("(")[0].strip()


def _no_krivet(s):
    """내부 보고서용: '(직능연)' '직능연 기준' 등 출처 표기를 지운다."""
    t = re.sub(r"[\(\[〔（][^\)\]〕）]*직능연[^\)\]〕）]*[\)\]〕）]", "", str(s or ""))
    t = re.sub(r"직능연\s*(기준|출처)?\s*[:：]?", "", t)
    return re.sub(r"\s{2,}", " ", t).strip(" ,·/|-")


def _cut_sent(s, limit):
    """limit 안에서 문장 경계('다.' 등)로 자른다 — 중간 절단 방지."""
    s = str(s or "").strip()
    if len(s) <= limit:
        return s
    cut = s[:limit]
    for mark in ("다.", "."):
        i = cut.rfind(mark)
        if i >= limit * 0.5:
            return cut[:i + len(mark)]
    return cut.rstrip() + "…"


def make_pref_foreword(preferred, target_month):
    """제2부 우대 총평 — 제1부 총평과 대구 (LLM 실패 시 통계 문장 폴백)"""
    if not preferred:
        return ""
    dist = Counter(_no_krivet(r.get("우대분류", "")) or "기타" for r in preferred)
    risky = [r.get("법령명", "") for r in preferred
             if _code(r.get("Track1_위험도", "")) in ("C", "H")]
    fallback = (f"이달 국가기술자격 우대 조항 신설·변경은 총 {len(preferred)}건으로, "
                + ", ".join(f"{k} {v}건" for k, v in dist.most_common()) + "이 확인되었다."
                + (f" 이 중 {len(risky)}건은 경력이음형 자격제도와의 정합성 검토가 필요한 임계·고위험 유형이다." if risky else ""))
    prompt = f"""당신은 한국산업인력공단의 자격정책 분석관입니다. {target_month[:4]}년 {int(target_month[4:6])}월
국가기술자격 취득자 우대 조항 신설·변경 현황을 보고서 문체 2~3문장으로 총평하세요.
- 총 {len(preferred)}건 / 우대분류 분포: {dict(dist.most_common())}
- 경력이음 정합성 주의(임계 C·고위험 H): {len(risky)}건 {risky[:3]}
과장 없이 담백하게, 문장만 출력하세요."""
    try:
        out = _clean_citations(_ask(prompt)).strip()
        return out if out else fallback
    except Exception:
        return fallback


def make_details(selected):
    """선별된 법령별 상세(파급효과/배경/내용/효과) 작성."""
    print(f"🧠 [2-3] 핵심 {len(selected)}건 상세 분석 작성 중...")
    results = []
    for idx, r in enumerate(selected, 1):
        name = r.get("법령명", "")
        print(f"   📝 [{idx}/{len(selected)}] {name[:30]} ...", end=" ", flush=True)
        prompt = DETAIL_PROMPT.format(
            persona=PERSONA, law_name=name,
            enf_date=norm_date(r.get("시행일자", "")),
            dept=r.get("소관부처", ""),
            certs=r.get("관련 종목", ""),
            summary=r.get("주요 제·개정내용", ""),
            util_detail=r.get("활용도_상세", ""),
        )
        raw = _ask_with_search(prompt, temperature=0.2)
        data = _parse_detail_json(raw)
        # 파싱 실패(핵심 항목 공백) 시 재시도 — 타임아웃 아님, 재질의만.
        # 웹검색 응답이 JSON을 오염시키는 경우가 있어 2차는 검색 없이, 3차는 온도를 바꿔 시도.
        if not (data.get("bg") or data.get("main")):
            print("⚠️ 상세 JSON 해석 실패 → 재시도(검색 미사용)", end=" ", flush=True)
            data = _parse_detail_json(_ask(prompt, temperature=0.2))
        if not (data.get("bg") or data.get("main")):
            print("→ 재시도(응답 다양화)", end=" ", flush=True)
            data = _parse_detail_json(_ask(prompt, temperature=0.5))
        merged = dict(r)
        merged["impact_3lines"] = data.get("impact_3lines", [])
        merged["bg"] = data.get("bg", "")
        merged["main"] = data.get("main", "")
        merged["effect"] = data.get("effect", "")
        results.append(merged)
        print("✅" if (merged.get("bg") or merged.get("main")) else "⚠️ 상세 누락 — 프레임만 기록됨")
    return results


def _parse_detail_json(raw):
    """AI 응답에서 JSON 객체 추출 (펜스·줄바꿈·인용 정리 + 균형 중괄호·콤마 보정).
    웹검색 결합 응답은 JSON 앞뒤에 인용·잔여 텍스트가 붙기 쉬워, 탐욕적 {.*} 만으로는
    파싱이 깨질 수 있다 → ①첫 '{'부터 균형 잡힌 블록 ②탐욕 매칭 순으로 시도하고,
    각 후보에 트레일링 콤마 보정까지 적용한다."""
    if not raw:
        return {}
    s = raw.replace("```json", "").replace("```JSON", "").replace("```", "").strip()
    s = s.replace("\n", " ").replace("\r", " ")
    cands = []
    i = s.find("{")
    if i >= 0:                     # 후보1: 균형 중괄호 블록 (뒤쪽 잔여 텍스트에 강함)
        depth = 0
        for j in range(i, len(s)):
            if s[j] == "{":
                depth += 1
            elif s[j] == "}":
                depth -= 1
                if depth == 0:
                    cands.append(s[i:j + 1])
                    break
    m = re.search(r"\{.*\}", s, re.DOTALL)
    if m and (not cands or m.group(0) != cands[0]):   # 후보2: 기존 탐욕 매칭
        cands.append(m.group(0))
    for c in cands:
        for fix in (c, re.sub(r",\s*([}\]])", r"\1", c)):   # 트레일링 콤마 보정
            try:
                data = json.loads(fix)
                return {k: _clean_citations(v) for k, v in data.items()}
            except Exception:
                continue
    return {}


def _clean_citations(value):
    """문자열/리스트에서 [1], [2,3] 같은 각주 인용 표시를 제거."""
    def clean(s):
        if not isinstance(s, str):
            return s
        # [1], [1, 2], [출처] 등 대괄호 인용 제거
        s = re.sub(r"\s*\[\d+(?:\s*,\s*\d+)*\]", "", s)
        return s.strip()
    if isinstance(value, list):
        return [clean(x) for x in value]
    return clean(value)


# ============================================================
# [3-A] 차트 생성 (부처별 TOP 5)
# ============================================================
def _korean_font():
    """GitHub Actions(Ubuntu)에 설치된 한글 폰트 경로 자동 탐색."""
    candidates = [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJKkr-Regular.otf",
    ]
    for p in candidates:
        if os.path.exists(p):
            return font_manager.FontProperties(fname=p)
    return None  # 없으면 기본 폰트 (한글 깨질 수 있음 → yml에서 폰트 설치)


def make_chart(high, out_path):
    print("📊 [3] 부처별 차트 생성 중...")
    fm = _korean_font()
    if fm:
        plt.rcParams["font.family"] = fm.get_name()
    plt.rcParams["axes.unicode_minus"] = False

    dept_counts = Counter(r.get("소관부처", "기타") for r in high)
    top5 = dept_counts.most_common(5)
    names = [x[0] for x in top5]
    values = [x[1] for x in top5]

    fig, ax = plt.subplots(figsize=(7, 3.8))
    bars = ax.bar(names, values, color="#2E5A88", width=0.6)
    ax.set_ylabel("관련 법령 건수", fontproperties=fm, fontsize=11)
    ax.set_title("부처별 자격 관련 법령 제·개정 현황",
                 fontproperties=fm, fontsize=13, fontweight="bold", pad=12)
    for bar in bars:
        h = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2, h, f"{int(h)}건",
                ha="center", va="bottom", fontproperties=fm, fontsize=10, fontweight="bold")
    ax.set_xticks(range(len(names)))
    ax.set_xticklabels(names, fontproperties=fm, fontsize=9, rotation=10)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", linestyle="--", alpha=0.4)
    if values:
        ax.set_ylim(0, max(values) * 1.18)
    plt.tight_layout()
    plt.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close()


# 분야 분류용 키워드 (자격 종목명에 포함된 단어로 분야 판별)
FIELD_KEYWORDS = {
    "건설·건축": ["건축", "토목", "건설", "조경", "도시계획", "측량", "지적"],
    "전기·전자": ["전기", "전자", "반도체", "정보통신", "정보처리", "임베디드", "전파"],
    "기계": ["기계", "금형", "용접", "자동차", "차량", "항공", "조선", "철도차량"],
    "환경·에너지": ["환경", "대기", "수질", "폐기물", "에너지", "신재생", "가스", "원자력", "기상"],
    "화학·재료": ["화공", "화학", "금속", "재료", "세라믹", "섬유", "위험물"],
    "농림·수산": ["산림", "농화학", "수산", "축산", "시설원예", "종자", "어업", "조경"],
    "IT·디지털": ["정보관리", "컴퓨터", "빅데이터", "인공지능", "정보보안"],
    "안전·소방": ["안전", "소방", "방재", "재난", "비파괴"],
}


def classify_field(cert_string):
    """자격 종목 문자열에서 가장 많이 매칭되는 분야를 반환. 없으면 '기타'."""
    certs = str(cert_string)
    scores = {field: sum(certs.count(kw) for kw in kws)
              for field, kws in FIELD_KEYWORDS.items()}
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "기타"


def make_field_chart(high, out_path):
    """분야별 분포 가로 막대 차트 (자격 종목 키워드 기반)."""
    print("📊 [3-2] 분야별 분포 차트 생성 중...")
    fm = _korean_font()
    if fm:
        plt.rcParams["font.family"] = fm.get_name()
    plt.rcParams["axes.unicode_minus"] = False

    field_counts = Counter(classify_field(r.get("관련 종목", "")) for r in high)
    # 많은 순 정렬 (가로 막대는 아래→위라 역순으로)
    items = field_counts.most_common()
    labels = [x[0] for x in items][::-1]
    values = [x[1] for x in items][::-1]

    fig, ax = plt.subplots(figsize=(7, 3.8))
    bars = ax.barh(labels, values, color="#0F6E56", height=0.6)
    ax.set_xlabel("관련 법령 건수", fontproperties=fm, fontsize=11)
    ax.set_title("자격 분야별 관련 법령 분포",
                 fontproperties=fm, fontsize=13, fontweight="bold", pad=12)
    for bar in bars:
        w = bar.get_width()
        ax.text(w, bar.get_y() + bar.get_height() / 2, f" {int(w)}건",
                ha="left", va="center", fontproperties=fm, fontsize=9, fontweight="bold")
    ax.set_yticks(range(len(labels)))
    ax.set_yticklabels(labels, fontproperties=fm, fontsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="x", linestyle="--", alpha=0.4)
    if values:
        ax.set_xlim(0, max(values) * 1.15)
    plt.tight_layout()
    plt.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close()


# ============================================================
# [3-B] 이슈브리핑 docx 생성 — 우리가 디자인한 그대로
# ============================================================
def _set_cell_bg(cell, hex_color):
    """표 셀 배경색 (python-docx 기본 미지원 → XML 직접 주입)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _run(p, text, size=11, color=None, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def build_briefing_docx(target_month, total_laws, related_count,
                        big_increase, foreword, issues, chart_path, out_path,
                        field_chart_path=None, preferred=None, pref_foreword=""):
    preferred = preferred or []
    print("📄 [4-1] 이슈브리핑 docx 생성 중...")
    year, month = target_month[:4], str(int(target_month[4:6]))
    doc = Document()

    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(11)

    # 페이지: US Letter, 여백 약 2.5cm
    sec = doc.sections[0]
    sec.page_width = Cm(21.59)
    sec.page_height = Cm(27.94)
    sec.top_margin = Cm(2.3)
    sec.bottom_margin = Cm(2.3)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # --- 제목 ---
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "국가기술자격 관련 법령", size=20, color=NAVY, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Issue Briefing", size=26, color=BLUE, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"〈{year}년 {month}월호〉", size=13, color=GRAY, bold=True)
    _add_bottom_border(p, NAVY)

    # --- 목차 (한 권 2부 구성 안내) ---
    doc.add_paragraph()
    p = doc.add_paragraph(); _run(p, "■ 목차", size=12, color=BLUE, bold=True)
    _toc = [
        ("개요·모니터링 요약", "이달의 숫자를 한눈에"),
        ("제1부  자격 활용도 동향", f"핵심 이슈 {len(issues)}건 심층 분석 + 분포 차트"),
        ("제2부  자격 우대사항 동향", f"핵심 사례 {min(5, len(preferred))}건 심층 분석 (정책·구직자 관점)"),
        ("붙임  월간 상세목록(xlsx)", "총괄현황표 / 자격활용도분석 / 우대사항분석"),
    ]
    for _t, _d in _toc:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
        _run(p, f"· {_t}", size=10.5, color=NAVY, bold=True)
        _run(p, f"  —  {_d}", size=10, color=GRAY)

    # 전체 시행 법령 수 표시용 (총괄현황표에서 못 읽으면 None)
    has_total = bool(total_laws) and total_laws > related_count
    total_str = f"{total_laws}건" if has_total else None

    # --- 개요 ---
    p = doc.add_paragraph(); _run(p, "■ 개요", size=12, color=BLUE, bold=True)
    # 그달 마지막 날 계산 (1월=31, 2월=28/29 등)
    last_day = calendar.monthrange(int(year), int(month))[1]
    overview = [
        ("조사기간", f"{year}년 {month}월 1일 ~ {month}월 {last_day}일"),
    ]
    if has_total:
        overview.append(("조사대상", f"{month}월 시행 법령 총 {total_str} (국가기술자격 관련 {related_count}건)"))
    else:
        overview.append(("조사대상", f"{month}월 시행 국가기술자격 관련 법령 {related_count}건"))
    overview.append(("주요내용", f"자격 활용도 변동 {len(issues)}개 핵심 사례"))
    for label, val in overview:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
        _run(p, f"✓ {label} : ", size=10.5, color=NAVY, bold=True)
        _run(p, val, size=10.5)

    # --- 모니터링 요약 ---
    p = doc.add_paragraph(); _run(p, "■ 모니터링 요약", size=12, color=BLUE, bold=True)
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
    if has_total:
        # 전체 수를 아는 경우: "전체 N건 중, 관련 M건"
        _run(p, f"· {month}월 전체 시행 법령 ", size=10.5)
        _run(p, total_str, size=10.5, color=NAVY, bold=True)
        _run(p, " 중, 국가기술자격 관련 법령은 ", size=10.5)
        _run(p, f"{related_count}건", size=10.5, color=NAVY, bold=True)
        _run(p, "으로 조사", size=10.5)
    else:
        # 전체 수를 모르는 경우: 관련 법령 수만
        _run(p, f"· {month}월 국가기술자격 관련 법령은 ", size=10.5)
        _run(p, f"{related_count}건", size=10.5, color=NAVY, bold=True)
        _run(p, "으로 조사", size=10.5)
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
    _run(p, "· 그중 자격 활용도가 ", size=10.5)
    _run(p, "대폭 증가", size=10.5, color=RGBColor(0xC5, 0x5A, 0x11), bold=True)
    _run(p, f"한 법령은 {big_increase}건으로 조사", size=10.5)

    # --- 총평 (회색 박스) ---
    h = doc.add_paragraph(); _run(h, "이달의 주요 정책 트렌드 및 총평", size=13, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4); p.paragraph_format.right_indent = Cm(0.4)
    _run(p, foreword, size=10.5, color=DARKGRAY)
    _shade_paragraph(p, "F2F2F2")
    _add_left_border(p, BLUE)

    # --- 차트 (부처별 + 분야별) ---
    h = doc.add_paragraph(); _run(h, "데이터 시각화 분석", size=13, color=NAVY, bold=True)
    if os.path.exists(chart_path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(chart_path, width=Cm(12.5))
    if field_chart_path and os.path.exists(field_chart_path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(field_chart_path, width=Cm(12.5))

    # --- 페이지 나눔 → 요약표 ---
    doc.add_page_break()
    h = doc.add_paragraph(); _run(h, "〈 주요 제·개정 법령 요약 〉", size=12, color=NAVY, bold=True)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "제1부  자격 활용도 동향", size=15, color=NAVY, bold=True)
    _add_bottom_border(p, NAVY)
    _build_summary_table(doc, issues)

    # --- 페이지 나눔 → 법령별 상세 ---
    doc.add_page_break()
    h = doc.add_paragraph(); _run(h, "〈 핵심 법령 상세 분석 〉", size=12, color=NAVY, bold=True)
    for i, it in enumerate(issues, 1):
        _build_detail_card(doc, i, it)

    # 머리말/푸터
    _add_header_footer(doc)
    _build_pref_part(doc, preferred, pref_foreword)

    doc.save(out_path)


def _build_pref_part(doc, preferred, pref_foreword=""):
    """제2부: 자격 우대사항 동향 — 총평 + 분포 표 + 정책(T1)·구직자(T2) 관점 + 중처법 + 사례 카드"""
    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "제2부  자격 우대사항 동향", size=15, color=NAVY, bold=True)
    _add_bottom_border(p, NAVY)
    doc.add_paragraph()
    if not preferred:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
        _run(p, "· 이달 신설·변경된 자격 우대 조항은 없습니다.", size=10.5)
        return

    if pref_foreword:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.2)
        _shade_paragraph(p, "F2F6FB")
        _run(p, pref_foreword, size=10.5, color=DARKGRAY)
        doc.add_paragraph()

    # ── ① 분류체계 안내 — Track 1·2 읽는 법 (표 정리, brain 분석 기준 정의) ──
    p = doc.add_paragraph(); _run(p, "■ 분류체계 안내 — Track 1·2 읽는 법", size=12, color=BLUE, bold=True)

    def _guide_table(title, rows, note=None):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
        _run(p, title, size=10, color=NAVY, bold=True)
        t = doc.add_table(rows=1 + len(rows), cols=3)
        t.alignment = 1
        widths = (Cm(2.6), Cm(3.6), Cm(10.0))
        for j, h in enumerate(("구분", "코드·명칭", "의미")):
            c = t.rows[0].cells[j]
            _set_cell_bg(c, "1F3864")
            _run(c.paragraphs[0], h, size=9.5, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
        for i, (grp, code, desc) in enumerate(rows, 1):
            cells = t.rows[i].cells
            if grp:
                _run(cells[0].paragraphs[0], grp, size=9, color=NAVY, bold=True)
                _set_cell_bg(cells[0], "EDF2F9")
            _run(cells[1].paragraphs[0], code, size=9, color=NAVY, bold=True)
            _run(cells[2].paragraphs[0], desc, size=9, color=GRAY)
        for r_ in t.rows:
            for j, w in enumerate(widths):
                r_.cells[j].width = w
        if note:
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
            _run(p, note, size=8.5, color=GRAY)
        doc.add_paragraph()

    _guide_table("· Track 1 (정책 관점) — 법령이 자격을 '어떻게' 다루는지", [
        ("취급유형", "A 신분형성형", "자격 취득자만 해당 명칭·신분을 사용할 수 있음"),
        ("", "B 영업요건형", "기업이 사업을 등록·지정받으려면 자격자 고용이 요건"),
        ("", "C 직역독점형", "특정 업무·행위는 자격자만 수행 가능"),
        ("", "D 인사가산형", "채용·보수·승진 등에서 가점 부여"),
        ("", "E 검정연계형", "다른 시험의 응시자격 부여 또는 과목 면제"),
        ("위험도", "C 임계위험", "오직 단일 자격만 인정하고 대체 경로가 전혀 없음"),
        ("", "H 고위험", "'자격 + 경력 N년'을 동시에 요구"),
        ("", "M 중위험", "복수의 자격을 OR 조건으로 대체 가능"),
        ("", "L 저위험", "관련 학과 졸업 + 경력 등으로 진입 우회 가능"),
        ("", "N 무관", "직역 진입을 막지 않는 단순 부가우대"),
    ], note="※ 위험도 = 선경력을 요구하는 '경력이음형' 자격제도와 충돌하는 정도")
    _guide_table("· Track 2 (구직자 관점) — 취득자에게 '무엇이' 생기는지", [
        ("Ⅰ 직업창출", "Ⅰ-1 면허전환형", "자격 취득이 행정청 면허 발급으로 이어져 평생 직업·신분 부여"),
        ("", "Ⅰ-2 개업창업형", "자격자 본인이 단독으로 직무 수행·서명 가능 — 1인 사업 가능"),
        ("Ⅱ 취업관문", "Ⅱ-1 등록필수형", "사업체 등록·허가 시 자격자 보유가 의무"),
        ("", "Ⅱ-2 지정인력형", "검사·검정·인증·진단 등 국가 지정기관의 인력 요건"),
        ("", "Ⅱ-3 전속배치형", "해당 자격자만 선임 가능 (대체 불가)"),
        ("", "Ⅱ-4 선택배치형", "법령이 인정하는 복수 자격 중 택일 선임"),
        ("", "Ⅱ-5 현장배치형", "공사·사업장 규모·종별에 따라 현장 단위 배치 의무"),
        ("Ⅲ 부가우대", "Ⅲ-1 시험면제", "다른 자격·시험의 응시요건 완화나 과목 면제"),
        ("", "Ⅲ-2 인사우대", "채용·승진·보수에서의 가점·우대"),
        ("", "Ⅲ-3 위촉·자문", "위원 위촉, 자문·심의 참여 자격 부여"),
    ])

    p = doc.add_paragraph(); _run(p, "■ 이달의 우대사항 현황", size=12, color=BLUE, bold=True)
    dist = Counter(_no_krivet(r.get("우대분류", "")) or "기타" for r in preferred)
    rep = {}
    for r in preferred:
        rep.setdefault(_no_krivet(r.get("우대분류", "")) or "기타", r.get("법령명", ""))
    tbl = doc.add_table(rows=1 + len(dist), cols=3)
    tbl.alignment = 1
    for j, h in enumerate(("우대분류", "건수", "대표 법령")):
        cell = tbl.rows[0].cells[j]
        _set_cell_bg(cell, "1F3864")
        _run(cell.paragraphs[0], h, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
    for i, (k, v) in enumerate(dist.most_common(), 1):
        row = tbl.rows[i].cells
        _run(row[0].paragraphs[0], k, size=10, bold=True)
        _run(row[1].paragraphs[0], f"{v}건", size=10, color=NAVY, bold=True)
        _run(row[2].paragraphs[0], str(rep.get(k, ""))[:34], size=9.5)
    doc.add_paragraph()

    p = doc.add_paragraph(); _run(p, "■ 분류체계 현황 — 정책 관점 (Track 1)", size=12, color=BLUE, bold=True)
    ttype = Counter(_code(r.get("Track1_취급유형", "")) or "-" for r in preferred)
    tname = {"A": "신분형성", "B": "영업요건", "C": "직역독점", "D": "인사가산", "E": "검정연계"}
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
    _run(p, "· 취급유형 분포 : ", size=10.5)
    _run(p, " · ".join(f"{k} {tname[k]} {ttype.get(k, 0)}건" for k in "ABCDE" if ttype.get(k, 0)) or "-",
         size=10.5, color=NAVY, bold=True)
    risk = Counter(_code(r.get("Track1_위험도", "")) or "-" for r in preferred)
    name = {"C": "임계", "H": "고위험", "M": "중위험", "L": "저위험", "N": "무관"}
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
    _run(p, "· 위험도 분포 : ", size=10.5)
    _run(p, " · ".join(f"{k} {name[k]} {risk.get(k, 0)}건" for k in ("C", "H", "M", "L", "N") if risk.get(k, 0)) or "-",
         size=10.5, color=NAVY, bold=True)
    risky = [r for r in preferred if _code(r.get("Track1_위험도", "")) in ("C", "H")]
    if risky:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
        _run(p, "⚠ 정합성 검토 필요(임계·고위험) : ", size=10, color=RGBColor(0xB0, 0x30, 0x30), bold=True)
        _run(p, ", ".join(str(r.get("법령명", ""))[:24] for r in risky[:4])
             + (f" 외 {len(risky) - 4}건" if len(risky) > 4 else ""), size=10)

    p = doc.add_paragraph(); _run(p, "■ 분류체계 현황 — 구직자 관점 (Track 2)", size=12, color=BLUE, bold=True)
    grp = Counter()
    for r in preferred:
        c = _code(r.get("Track2_효용코드", ""))
        if c.startswith("Ⅰ"): grp["직업창출(Ⅰ)"] += 1
        elif c.startswith("Ⅱ"): grp["취업관문(Ⅱ)"] += 1
        elif c.startswith("Ⅲ"): grp["부가우대(Ⅲ)"] += 1
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
    _run(p, "· 효용 유형 : ", size=10.5)
    _run(p, " · ".join(f"{k} {v}건" for k, v in grp.most_common()) or "-", size=10.5, color=NAVY, bold=True)
    cert_cnt = Counter()
    for r in preferred:
        for c in str(r.get("관련 종목", "")).split(","):
            c = c.strip()
            if c and c != "없음":
                cert_cnt[c] += 1
    if cert_cnt:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
        _run(p, "· 우대 조항 다빈도 자격 : ", size=10.5)
        _run(p, ", ".join(f"{k}({v})" for k, v in cert_cnt.most_common(5)), size=10.5, color=NAVY, bold=True)

    sap = [r for r in preferred if str(r.get("중처법대상", "")).strip() == "대상"]
    if sap:
        p = doc.add_paragraph(); _run(p, "■ 중대재해처벌법 연계", size=12, color=BLUE, bold=True)
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
        _run(p, f"· 안전관리 의무와 직결된 우대 {len(sap)}건 : ", size=10.5)
        _run(p, ", ".join(str(r.get("법령명", ""))[:24] for r in sap[:3])
             + (f" 외 {len(sap) - 3}건" if len(sap) > 3 else ""), size=10.5, color=NAVY, bold=True)

    doc.add_paragraph()
    p = doc.add_paragraph(); _run(p, "■ 주요 사례", size=12, color=BLUE, bold=True)
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
    _run(p, "· 선정 기준 : 위험도 상위(임계 C → 고위험 H → …) — 정책 검토 우선순위 순", size=9, color=GRAY)
    _RORD = {"C": 0, "H": 1, "M": 2, "L": 3, "N": 4}
    _PORD = {"의무고용": 0, "직무권한부여": 1, "인사우대": 2, "시험면제": 3}
    top = sorted(preferred, key=lambda x: (_RORD.get(_code(x.get("Track1_위험도", "")), 9),
                                           _PORD.get(_no_krivet(x.get("우대분류", "")), 9),
                                           str(x.get("법령명", ""))))[:5]
    for i, r in enumerate(top, 1):
        doc.add_paragraph()
        p = doc.add_paragraph(); _add_left_border(p, NAVY)
        _run(p, f"  {i}. {r.get('법령명', '')}", size=11.5, color=NAVY, bold=True)
        if str(r.get("중처법대상", "")).strip() == "대상":
            _run(p, "   ⚠ 중처법 연계", size=9.5, color=RGBColor(0xB0, 0x30, 0x30), bold=True)
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.6)
        _run(p, f"[{_no_krivet(r.get('우대분류', '')) or '기타'}] ", size=10, color=BLUE, bold=True)
        _run(p, _summarize_certs(r.get("관련 종목", "")), size=10)
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.6)
        _run(p, f"취급 {r.get('Track1_취급유형', '')} · 위험 {r.get('Track1_위험도', '')} · 효용 {r.get('Track2_효용코드', '')}",
             size=9, color=GRAY)
        summ = str(r.get("조문 요약", "")).strip()
        if summ:
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.6)
            _run(p, _cut_sent(summ, 700), size=9.5, color=DARKGRAY)
        det = str(r.get("상세 분석 결과", "")).strip()
        if det:
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.6)
            _run(p, "→ " + _cut_sent(det, 500), size=9, color=GRAY, italic=True)


def _build_summary_table(doc, issues):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    widths = [Cm(1.2), Cm(5.5), Cm(2.0), Cm(4.0), Cm(3.5)]
    hdr = table.rows[0].cells
    for c, (txt, w) in enumerate(zip(["연번", "법령명", "시행일", "관련 자격", "핵심 내용"], widths)):
        hdr[c].width = w
        _set_cell_bg(hdr[c], "1F3864")
        para = hdr[c].paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, txt, size=9, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
    for i, it in enumerate(issues, 1):
        cells = table.add_row().cells
        date = norm_date(it.get("시행일자", ""))
        date_fmt = f"{date[:4]}. {date[4:6]}. {date[6:8]}." if len(date) == 8 else date
        certs = _summarize_certs(it.get("관련 종목", ""))
        summ = str(it.get("주요 제·개정내용", ""))[:40] + "…"
        vals = [str(i), it.get("법령명", ""), date_fmt, certs, summ]
        aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT]
        for c, (v, w, al) in enumerate(zip(vals, widths, aligns)):
            cells[c].width = w
            para = cells[c].paragraphs[0]; para.alignment = al
            _run(para, v, size=8.5)


def _build_detail_card(doc, i, it):
    # 제목 바 (네이비)
    p = doc.add_paragraph()
    _run(p, f"{i}. {it.get('법령명','')}", size=11, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
    _shade_paragraph(p, "1F3864")
    # 메타 (연파랑)
    date = norm_date(it.get("시행일자", ""))
    date_fmt = f"{date[:4]}. {date[4:6]}. {date[6:8]}." if len(date) == 8 else date
    certs = _summarize_certs(it.get("관련 종목", ""))
    p = doc.add_paragraph()
    _run(p, "시행일 ", size=8.5, color=NAVY, bold=True); _run(p, f"{date_fmt}    ", size=8.5)
    _run(p, "관련 자격 ", size=8.5, color=NAVY, bold=True); _run(p, certs, size=8.5)
    _shade_paragraph(p, "D6E2F0")
    # 파급효과
    p = doc.add_paragraph(); _run(p, "▣ 자격증 파급효과", size=10, color=BLUE, bold=True)
    reason = str(it.get("_선정사유", "")).strip()
    if reason:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.2)
        _run(p, "선정 사유  ", size=8.5, color=NAVY, bold=True)
        _run(p, reason, size=8.5, color=GRAY)
    for k, line in enumerate(it.get("impact_3lines", []), 1):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.6)
        _run(p, f"{k}. ", size=9.5, color=NAVY, bold=True); _run(p, line, size=9.5)
    # 상세 분석
    p = doc.add_paragraph(); _run(p, "▣ 상세 분석", size=10, color=BLUE, bold=True)
    for lbl, val in [("□ 추진배경", it.get("bg", "")),
                     ("□ 주요 개정내용", it.get("main", "")),
                     ("□ 자격증 기대효과", it.get("effect", ""))]:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
        _run(p, lbl, size=9.5, color=DARKGRAY, bold=True)
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8)
        _run(p, val, size=9.5, color=DARKGRAY)
    # 개정 조문 (회색, 본문과 분리)
    p = doc.add_paragraph(); _run(p, "▣ 개정 조문", size=10, color=BLUE, bold=True)
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
    jomun = str(it.get("근거조문", "") or "(해당 없음)")
    _run(p, jomun, size=8.5, color=GRAY, italic=True)
    _shade_paragraph(p, "F2F2F2")
    # 구분선
    p = doc.add_paragraph(); _add_bottom_border(p, RGBColor(0xBF, 0xBF, 0xBF), dashed=True)


def _summarize_certs(cert_string):
    certs = [c.strip() for c in str(cert_string).split(",") if c.strip()]
    if len(certs) > 1:
        return f"{certs[0]} 등 {len(certs)}개 종목"
    return certs[0] if certs else "해당 없음"


# --- 단락 음영/테두리 유틸 (XML 직접 조작) ---
def _shade_paragraph(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _add_bottom_border(p, color, dashed=False):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "dashed" if dashed else "single")
    bottom.set(qn("w:sz"), "6"); bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    pbdr.append(bottom); pPr.append(pbdr)


def _add_left_border(p, color):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "24"); left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    pbdr.append(left); pPr.append(pbdr)


def _add_header_footer(doc):
    sec = doc.sections[0]
    hdr = sec.header.paragraphs[0]; hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(hdr, "국가기술자격 관련 법령 Issue Briefing", size=7.5, color=GRAY)
    ftr = sec.footer.paragraphs[0]; ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(ftr, "한국산업인력공단", size=7.5, color=GRAY)


# ============================================================
# [3-C] 모니터링 결과 xlsx 생성 — 근거조문 + 하이퍼링크
# ============================================================
def build_monitor_xlsx(target_month, total_laws, high, simple, out_path, preferred=None):
    preferred = preferred or []
    print("📊 [4-2] 모니터링 결과 xlsx 생성 중...")
    year, month = target_month[:4], str(int(target_month[4:6]))
    YM = f"{year}년 {month}월"

    XLNAVY = "1F3864"; XLBLUE = "2E5A88"; XLLIGHTBLUE = "D6E2F0"
    YELLOW = "FFF2CC"; LINKBLUE = "0563C1"; HGRAY = "404040"
    thin = Side(style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    def hdr(cell, fill=XLNAVY, color="FFFFFF"):
        cell.font = XLFont(name=FONT, bold=True, color=color, size=10)
        cell.fill = PatternFill("solid", fgColor=fill); cell.alignment = center; cell.border = border

    def body(cell, fill=None, align=left, size=9):
        cell.font = XLFont(name=FONT, size=size); cell.alignment = align; cell.border = border
        if fill:
            cell.fill = PatternFill("solid", fgColor=fill)

    def law_url(nm):
        return f"https://www.law.go.kr/법령/{nm}"

    wb = Workbook()

    # 시트1: 요약
    ws = wb.active; ws.title = "총괄현황표"
    ws.merge_cells("A1:C1")
    ws["A1"] = f"{YM} 국가기술자격 관련 법령 제·개정사항 모니터링 결과"
    ws["A1"].font = XLFont(name=FONT, bold=True, size=14, color=XLNAVY)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30
    related_total = len(high) + len(simple)
    has_total = bool(total_laws) and total_laws > related_total
    ws.merge_cells("A2:C2")
    total_label = f"총 {total_laws}건" if has_total else f"국가기술자격 관련 {related_total}건"
    ws["A2"] = f"※ {YM} 시행 법령 : {total_label}"
    ws["A2"].font = XLFont(name=FONT, size=10, color=HGRAY)

    util = Counter(r.get("활용도_구분", "") for r in high)
    stats = [
        ("구분", "건수", "비고", True),
        ("총 시행 법령", total_laws if has_total else "-", "당월 시행된 전체 법령", False),
        ("국가기술자격 관계 법령", related_total, "자격 관련 있는 법령 (아래 합계)", False),
        ("  ① 활용·관련 높은 법령", len(high), "자격 활용도에 영향", False),
        ("      · 대폭 증가", util.get("대폭 증가", 0), "자격 수요 크게 증가", False),
        ("      · 소폭 증가", util.get("소폭 증가", 0), "자격 수요 소폭 증가", False),
        ("      · 현상 유지", util.get("현상 유지", 0), "변동 미미", False),
        ("  ② 단순 관련 법령", len(simple), "자격 언급되나 활용도 변동 적음", False),
        ("  ③ 자격 우대 신설·변경", len(preferred), "우대여부 O — 취득자 우대 조항", False),
    ]
    for i, (label, cnt, note, is_h) in enumerate(stats):
        r = 4 + i
        ws[f"A{r}"], ws[f"B{r}"], ws[f"C{r}"] = label, cnt, note
        if is_h:
            for col in "ABC":
                hdr(ws[f"{col}{r}"])
        else:
            fill = XLLIGHTBLUE if label == "국가기술자격 관계 법령" else None
            body(ws[f"A{r}"], fill=fill, align=left, size=10)
            body(ws[f"B{r}"], fill=fill or (YELLOW if "①" in label else None), align=center, size=10)
            body(ws[f"C{r}"], fill=fill, align=left, size=9)
            ws[f"A{r}"].font = XLFont(name=FONT, size=10,
                                      bold=("관계 법령" in label or "①" in label or "②" in label))
            ws[f"B{r}"].font = XLFont(name=FONT, size=10, bold=True, color=XLNAVY)
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 38

    # 시트2: 자격활용도분석 — 연관높음 + 단순관련 통합 (일일 xlsx와 같은 3탭 체계)
    ws2 = wb.create_sheet("자격활용도분석")
    ws2.merge_cells("A1:J1")
    ws2["A1"] = f"1. 자격 활용도 분석 : 연관높음 {len(high)}건 + 단순관련 {len(simple)}건"
    ws2["A1"].font = XLFont(name=FONT, bold=True, size=12, color=XLNAVY)
    ws2.row_dimensions[1].height = 24
    h_head = ["연번", "연관도", "법령명", "시행일자", "소관부처", "주요 제·개정내용",
              "관련 종목", "활용도 구분", "활용도 상세", "근거 조문"]
    h_w = [6, 10, 28, 12, 14, 42, 34, 11, 40, 34]
    for c, (t, w) in enumerate(zip(h_head, h_w), 1):
        hdr(ws2.cell(row=2, column=c, value=t))
        ws2.column_dimensions[get_column_letter(c)].width = w
    order = {"대폭 증가": 0, "대폭 감소": 1, "소폭 증가": 2, "소폭 감소": 3, "현상 유지": 4}
    rows_all = sorted(high, key=lambda x: order.get(x.get("활용도_구분", ""), 9)) + list(simple)
    for i, r in enumerate(rows_all, 1):
        row = 2 + i
        d = norm_date(r.get("시행일자", ""))
        df = f"{d[:4]}-{d[4:6]}-{d[6:8]}" if len(d) == 8 else d
        rel = str(r.get("연관도", "")).strip()
        is_big = r.get("활용도_구분", "") in ("대폭 증가", "대폭 감소")
        fill = YELLOW if is_big else None
        body(ws2.cell(row=row, column=1, value=i), fill=fill, align=center)
        body(ws2.cell(row=row, column=2, value=rel), fill=fill, align=center)
        c3 = ws2.cell(row=row, column=3, value=r.get("법령명", ""))
        c3.hyperlink = law_url(r.get("법령명", ""))
        c3.font = XLFont(name=FONT, size=9, color=LINKBLUE, underline="single")
        c3.alignment = left; c3.border = border
        if fill:
            c3.fill = PatternFill("solid", fgColor=fill)
        body(ws2.cell(row=row, column=4, value=df), fill=fill, align=center)
        body(ws2.cell(row=row, column=5, value=r.get("소관부처", "")), fill=fill, align=left)
        body(ws2.cell(row=row, column=6, value=str(r.get("주요 제·개정내용", ""))[:200]), fill=fill, align=left)
        body(ws2.cell(row=row, column=7, value=str(r.get("관련 종목", ""))[:150]), fill=fill, align=left)
        body(ws2.cell(row=row, column=8, value=r.get("활용도_구분", "")), fill=fill, align=center)
        body(ws2.cell(row=row, column=9, value=str(r.get("활용도_상세", ""))[:250]), fill=fill, align=left)
        body(ws2.cell(row=row, column=10, value=str(r.get("근거조문", ""))[:200]), fill=fill, align=left)
        ws2.row_dimensions[row].height = 48
    ws2.freeze_panes = "A3"

    # 시트3: 우대사항분석 — 우대여부=O (RADAR 혈통 관점)
    ws3 = wb.create_sheet("우대사항분석")
    ws3.merge_cells("A1:J1")
    ws3["A1"] = f"2. 자격 우대사항 분석 : 우대 신설·변경 {len(preferred)}건"
    ws3["A1"].font = XLFont(name=FONT, bold=True, size=12, color=XLNAVY)
    ws3.row_dimensions[1].height = 24
    p_head = ["연번", "법령명", "시행일자", "우대분류", "관련 종목", "취급유형", "위험도",
              "효용코드", "중처법", "조문 요약"]
    p_w = [6, 30, 12, 12, 32, 18, 14, 18, 9, 48]
    for c, (t, w) in enumerate(zip(p_head, p_w), 1):
        hdr(ws3.cell(row=2, column=c, value=t), fill=XLBLUE)
        ws3.column_dimensions[get_column_letter(c)].width = w

    _RORD = {"C": 0, "H": 1, "M": 2, "L": 3, "N": 4}
    _PORD = {"의무고용": 0, "직무권한부여": 1, "인사우대": 2, "시험면제": 3}
    for i, r in enumerate(sorted(preferred,
                                 key=lambda x: (_RORD.get(_code(x.get("Track1_위험도", "")), 9),
                                                _PORD.get(_no_krivet(x.get("우대분류", "")), 9),
                                                str(x.get("법령명", "")))), 1):
        row = 2 + i
        d = norm_date(r.get("시행일자", ""))
        df = f"{d[:4]}-{d[4:6]}-{d[6:8]}" if len(d) == 8 else d
        body(ws3.cell(row=row, column=1, value=i), align=center)
        c2 = ws3.cell(row=row, column=2, value=r.get("법령명", ""))
        c2.hyperlink = law_url(r.get("법령명", ""))
        c2.font = XLFont(name=FONT, size=9, color=LINKBLUE, underline="single")
        c2.alignment = left; c2.border = border
        body(ws3.cell(row=row, column=3, value=df), align=center)
        body(ws3.cell(row=row, column=4, value=_no_krivet(r.get("우대분류", ""))), align=center)
        body(ws3.cell(row=row, column=5, value=str(r.get("관련 종목", ""))[:150]), align=left)
        body(ws3.cell(row=row, column=6, value=r.get("Track1_취급유형", "")), align=center)
        body(ws3.cell(row=row, column=7, value=r.get("Track1_위험도", "")), align=center)
        body(ws3.cell(row=row, column=8, value=r.get("Track2_효용코드", "")), align=center)
        body(ws3.cell(row=row, column=9, value=r.get("중처법대상", "")), align=center)
        body(ws3.cell(row=row, column=10, value=_cut_sent(r.get("조문 요약", ""), 400)), align=left)
        ws3.row_dimensions[row].height = 44
    ws3.freeze_panes = "A3"

    wb.save(out_path)


# ============================================================
# [4] 웹훅으로 두 파일 첨부 발송 (Make.com → 메일)
# ============================================================
def send_via_webhook(target_month, docx_path, xlsx_path, stats, pdf_path=None):
    if not WEBHOOK_URL:
        print("ℹ️ BRIEFING_WEBHOOK_URL이 없어 메일 발송을 건너뜁니다. (로컬 테스트 모드)")
        return
    print("📧 [5] 웹훅으로 보고서 발송 중...")
    year, month = target_month[:4], str(int(target_month[4:6]))
    data = {
        "system": "HRDK Q-RADAR",
        "source": "qradar_briefing",
        "subject": f"[Q-RADAR 이슈브리핑] {year}년 {month}월호 국가기술자격 법령 동향",
        "month": f"{year}년 {month}월",
        "total": str(stats.get("total", "")),
        "related": str(stats.get("related", "")),
        "big": str(stats.get("big", "")),
        "preferred": str(stats.get("preferred", "")),
    }
    files = {}
    fh1 = open(docx_path, "rb"); fh2 = open(xlsx_path, "rb")
    fh3 = open(pdf_path, "rb") if (pdf_path and os.path.exists(pdf_path)) else None
    try:
        files = {
            "file1": (os.path.basename(docx_path), fh1,
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            "file2": (os.path.basename(xlsx_path), fh2,
                      "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
        }
        if fh3:
            files["file3"] = (os.path.basename(pdf_path), fh3, "application/pdf")
        requests.post(WEBHOOK_URL, data=data, files=files, timeout=60)
        print("   ✅ 발송 완료")
    except Exception as e:
        print(f"   ⚠️ 발송 실패: {e}")
    finally:
        fh1.close(); fh2.close()
        if fh3:
            fh3.close()


# ============================================================
# 메인
# ============================================================
def main():
    if not TARGET_MONTH or len(TARGET_MONTH) != 6 or not TARGET_MONTH.isdigit():
        raise SystemExit("❌ TARGET_MONTH는 6자리 숫자(YYYYMM)여야 합니다. 예: 202601")

    print(f"🚀 {TARGET_MONTH} 이슈브리핑 생성 시작\n" + "=" * 50)

    # 1) 데이터 (전체 시행 법령 수도 총괄현황표에서 함께 가져옴)
    high, simple, preferred, total_laws = fetch_month_data(TARGET_MONTH)
    if not high and not simple:
        raise SystemExit(f"❌ {TARGET_MONTH} 데이터가 시트에 없습니다.")

    related_count = len(high) + len(simple)
    util = Counter(r.get("활용도_구분", "") for r in high)
    big_increase = util.get("대폭 증가", 0) + util.get("대폭 감소", 0)

    # 2) AI: 선별 → 총평 → 상세
    big_laws = [r for r in high if r.get("활용도_구분", "") in ("대폭 증가", "대폭 감소")]
    selected = select_top_laws(big_laws, TOP_N)
    foreword = make_foreword(selected, TARGET_MONTH)
    issues = make_details(selected)
    pref_foreword = make_pref_foreword(preferred, TARGET_MONTH)

    # 3) 차트 + 두 산출물
    chart_path = "/tmp/chart.png"
    make_chart(high, chart_path)
    field_chart_path = "/tmp/field_chart.png"
    make_field_chart(high, field_chart_path)

    docx_path = f"/tmp/이슈브리핑_{TARGET_MONTH}.docx"
    xlsx_path = f"/tmp/모니터링결과_{TARGET_MONTH}.xlsx"
    build_briefing_docx(TARGET_MONTH, total_laws, related_count, big_increase,
                        foreword, issues, chart_path, docx_path,
                        field_chart_path=field_chart_path, preferred=preferred, pref_foreword=pref_foreword)
    build_monitor_xlsx(TARGET_MONTH, total_laws, high, simple, xlsx_path, preferred=preferred)

    pdf_path = f"/tmp/이슈브리핑_{TARGET_MONTH}.pdf"
    try:
        from briefing_pdf import build_briefing_pdf
        build_briefing_pdf(TARGET_MONTH, total_laws, related_count, len(high), len(simple),
                           foreword, issues, preferred, pref_foreword,
                           chart_path, field_chart_path, pdf_path)
    except Exception as e:
        print(f"⚠️ PDF 디자인판 생성 건너뜀({e}) — docx·xlsx는 정상 진행")
        pdf_path = None

    # GitHub Actions가 가져갈 수 있게 현재 폴더에도 복사
    import shutil
    out_dir = os.environ.get("OUTPUT_DIR", ".")
    os.makedirs(out_dir, exist_ok=True)
    final_docx = os.path.join(out_dir, os.path.basename(docx_path))
    final_xlsx = os.path.join(out_dir, os.path.basename(xlsx_path))
    shutil.copy(docx_path, final_docx)
    shutil.copy(xlsx_path, final_xlsx)
    if pdf_path and os.path.exists(pdf_path):
        shutil.copy(pdf_path, os.path.join(out_dir, os.path.basename(pdf_path)))
    print(f"\n📁 생성 완료: {final_docx}, {final_xlsx}")

    # 4) 발송
    send_via_webhook(TARGET_MONTH, final_docx, final_xlsx,
                     {"total": total_laws or "-", "related": related_count, "big": big_increase, "preferred": len(preferred)}, pdf_path=pdf_path)

    print("=" * 50 + "\n✨ 이슈브리핑 생성 완료!")


if __name__ == "__main__":
    main()
